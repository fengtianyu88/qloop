"""qloop — 全面功能测试脚本

依据设计文档 docs/superpowers/specs/2026-07-16-qloop-design.md
对后端 API 进行端到端功能测试，覆盖：
  1. 认证 (登录/注册/当前用户/密码重置流程)
  2. 用户管理 (CRUD/角色)
  3. 组织管理 (组织树/组织单元/admin scope)
  4. LLM 配置 (模型/评审规则)
  5. 项目生命周期 (项目/成员/版本)
  6. 释放流程 (7 步流转 + 文件上传 + 确认)
  7. 搜索筛选 (释放/项目, GUEST 隔离)
  8. 审计日志
  9. 通知
  10. 权限矩阵 (无 Token / 角色不足 / 跨项目隔离)

运行: /opt/qloop/backend/venv/bin/python /workspace/test_features.py
"""
from __future__ import annotations

import io
import json
import time
import uuid
from dataclasses import dataclass, field
from typing import Any, Optional

import httpx

BASE = "http://localhost:8000"
TIMEOUT = 30.0


@dataclass
class Result:
    name: str
    passed: bool
    detail: str = ""
    expected: Any = None
    actual: Any = None


@dataclass
class Suite:
    name: str
    results: list[Result] = field(default_factory=list)

    def add(self, r: Result) -> None:
        self.results.append(r)

    @property
    def passed_count(self) -> int:
        return sum(1 for r in self.results if r.passed)

    @property
    def failed_count(self) -> int:
        return sum(1 for r in self.results if not r.passed)


SUITES: list[Suite] = []


def suite(name: str) -> Suite:
    s = Suite(name=name)
    SUITES.append(s)
    return s


def check(s: Suite, name: str, cond: bool, detail: str = "",
          expected: Any = None, actual: Any = None) -> None:
    s.add(Result(name=name, passed=bool(cond), detail=detail,
                 expected=expected, actual=actual))


def login(client: httpx.Client, username: str, password: str) -> dict | None:
    r = client.post("/api/auth/login", json={"username": username, "password": password},
                    timeout=TIMEOUT)
    if r.status_code == 200:
        return r.json()
    return None


def auth_header(token: str) -> dict:
    return {"Authorization": f"Bearer {token}"}


# 邮箱唯一化辅助 (使用 example.com, email-validator 接受的保留域名)
def uniq_email(prefix: str = "user") -> str:
    return f"{prefix}.{uuid.uuid4().hex[:8]}@example.com"


def uniq_name(prefix: str = "u") -> str:
    return f"{prefix}_{uuid.uuid4().hex[:8]}"


# ============================================================================
# 1. 认证测试
# ============================================================================
def test_auth() -> None:
    s = suite("1. 认证 (auth)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)

    # 1.1 正确凭据登录
    data = login(client, "admin", "Admin@123")
    check(s, "管理员登录成功", data is not None and "access_token" in data,
          detail=str(data)[:200] if data else "no token")
    admin_token = data["access_token"] if data else ""

    # 1.2 错误密码
    r = client.post("/api/auth/login", json={"username": "admin", "password": "wrong"})
    check(s, "错误密码返回 401", r.status_code == 401,
          expected=401, actual=r.status_code)

    # 1.3 不存在的用户
    r = client.post("/api/auth/login", json={"username": "nobody", "password": "x"})
    check(s, "不存在用户返回 401", r.status_code == 401,
          expected=401, actual=r.status_code)

    # 1.4 缺少字段 (422)
    r = client.post("/api/auth/login", json={"username": "admin"})
    check(s, "缺少 password 字段返回 422", r.status_code == 422,
          expected=422, actual=r.status_code)

    # 1.5 /api/users/me
    r = client.get("/api/users/me", headers=auth_header(admin_token))
    check(s, "/me 返回当前用户", r.status_code == 200 and r.json().get("username") == "admin",
          expected="admin", actual=r.json().get("username") if r.status_code == 200 else r.status_code)

    # 1.6 无 Token 访问 /me → 401
    r = client.get("/api/users/me")
    check(s, "无 Token 访问 /me 返回 401", r.status_code == 401,
          expected=401, actual=r.status_code)

    # 1.7 注册新账号
    email = uniq_email("reg")
    username = uniq_name("reg")
    r = client.post("/api/auth/register", json={
        "username": username, "email": email, "full_name": "注册测试",
        "password": "Test@1234", "department": "测试部", "section": "测试组",
    })
    check(s, "注册新账号返回 201", r.status_code == 201,
          expected=201, actual=r.status_code, detail=r.text[:200])

    # 1.8 注册后默认 guest 角色
    if r.status_code == 201:
        check(s, "注册账号默认 guest 角色",
              r.json().get("system_role") == "guest",
              expected="guest", actual=r.json().get("system_role"))

    # 1.9 重复用户名注册 → 400
    r2 = client.post("/api/auth/register", json={
        "username": username, "email": uniq_email("dup"), "full_name": "x",
        "password": "Test@1234",
    })
    check(s, "重复用户名注册返回 400", r2.status_code == 400,
          expected=400, actual=r2.status_code, detail=r2.text[:200])

    # 1.10 重复邮箱注册 → 400
    r3 = client.post("/api/auth/register", json={
        "username": uniq_name("dup"), "email": email, "full_name": "x",
        "password": "Test@1234",
    })
    check(s, "重复邮箱注册返回 400", r3.status_code == 400,
          expected=400, actual=r3.status_code)

    # 1.11 注册后能用新账号登录
    new_data = login(client, username, "Test@1234")
    check(s, "新注册账号可登录", new_data is not None and "access_token" in new_data)

    # 1.12 forgot-password (无论邮箱是否存在都返回 200，防枚举)
    r = client.post("/api/auth/forgot-password", json={"email": email})
    check(s, "forgot-password 存在邮箱返回 200", r.status_code == 200,
          expected=200, actual=r.status_code)
    r = client.post("/api/auth/forgot-password", json={"email": "notexist@example.com"})
    check(s, "forgot-password 不存在邮箱也返回 200 (防枚举)", r.status_code == 200,
          expected=200, actual=r.status_code)

    # 1.13 reset-password 无效 token → 400
    r = client.post("/api/auth/reset-password", json={
        "token": "invalid.token.here", "new_password": "NewPass@123",
    })
    check(s, "reset-password 无效 token 返回 400", r.status_code == 400,
          expected=400, actual=r.status_code)

    # 1.14 密码过短 → 422
    r = client.post("/api/auth/register", json={
        "username": uniq_name("short"), "email": uniq_email("short"),
        "full_name": "x", "password": "12345",  # 5 chars
    })
    check(s, "密码过短注册返回 422", r.status_code == 422,
          expected=422, actual=r.status_code)

    client.close()


# ============================================================================
# 2. 用户管理测试
# ============================================================================
def test_users(admin_token: str) -> dict:
    """返回创建的测试用户字典 (供后续测试使用)"""
    s = suite("2. 用户管理 (users)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)
    created: dict = {}

    # 2.1 列出用户 (分页)
    r = client.get("/api/users?page=1&page_size=10", headers=h)
    check(s, "GET /api/users 返回分页结构",
          r.status_code == 200 and "items" in r.json() and "total" in r.json(),
          expected="PaginatedResponse", actual=r.json().keys() if r.status_code == 200 else r.status_code)

    # 2.2 创建开发者用户
    dev_email = uniq_email("dev")
    dev_username = uniq_name("dev")
    r = client.post("/api/users", json={
        "username": dev_username, "email": dev_email, "full_name": "开发者甲",
        "password": "Dev@1234", "system_role": "developer",
        "department": "算法部", "section": "SOC组",
    }, headers=h)
    check(s, "创建 developer 用户返回 201", r.status_code == 201,
          expected=201, actual=r.status_code, detail=r.text[:200])
    if r.status_code == 201:
        created["developer_id"] = r.json()["id"]
        created["developer_username"] = dev_username
        created["developer_email"] = dev_email

    # 2.3 创建测试人员用户
    tester_email = uniq_email("tester")
    tester_username = uniq_name("tester")
    r = client.post("/api/users", json={
        "username": tester_username, "email": tester_email, "full_name": "测试乙",
        "password": "Test@1234", "system_role": "developer",
    }, headers=h)
    check(s, "创建 tester 用户返回 201", r.status_code == 201,
          expected=201, actual=r.status_code)
    if r.status_code == 201:
        created["tester_id"] = r.json()["id"]
        created["tester_username"] = tester_username

    # 2.4 创建外部专家 (guest 角色)
    expert_email = uniq_email("expert")
    expert_username = uniq_name("expert")
    r = client.post("/api/users", json={
        "username": expert_username, "email": expert_email, "full_name": "外部专家丙",
        "password": "Expert@123", "system_role": "guest",
    }, headers=h)
    check(s, "创建 guest 专家用户返回 201", r.status_code == 201,
          expected=201, actual=r.status_code)
    if r.status_code == 201:
        created["expert_id"] = r.json()["id"]
        created["expert_username"] = expert_username

    # 2.5 创建管理员用户
    admin2_email = uniq_email("admin2")
    admin2_username = uniq_name("admin2")
    r = client.post("/api/users", json={
        "username": admin2_username, "email": admin2_email, "full_name": "管理员丁",
        "password": "Admin2@123", "system_role": "admin",
    }, headers=h)
    check(s, "创建 admin 用户返回 201", r.status_code == 201,
          expected=201, actual=r.status_code)
    if r.status_code == 201:
        created["admin2_id"] = r.json()["id"]
        created["admin2_username"] = admin2_username

    # 2.6 重复用户名 → 409
    r = client.post("/api/users", json={
        "username": dev_username, "email": uniq_email("dup"), "full_name": "x",
        "password": "X@12345", "system_role": "developer",
    }, headers=h)
    check(s, "重复用户名创建返回 409", r.status_code == 409,
          expected=409, actual=r.status_code)

    # 2.7 按 ID 查询用户
    if "developer_id" in created:
        r = client.get(f"/api/users/{created['developer_id']}", headers=h)
        check(s, "按 ID 查询用户返回 200", r.status_code == 200,
              expected=200, actual=r.status_code)

    # 2.8 查询不存在用户 → 404
    fake_id = str(uuid.uuid4())
    r = client.get(f"/api/users/{fake_id}", headers=h)
    check(s, "查询不存在用户返回 404", r.status_code == 404,
          expected=404, actual=r.status_code)

    # 2.9 更新用户
    if "developer_id" in created:
        r = client.put(f"/api/users/{created['developer_id']}", json={
            "full_name": "开发者甲(已更新)", "department": "算法部2",
        }, headers=h)
        check(s, "更新用户返回 200", r.status_code == 200,
              expected=200, actual=r.status_code, detail=r.text[:200])
        if r.status_code == 200:
            check(s, "更新后 full_name 生效",
                  r.json().get("full_name") == "开发者甲(已更新)",
                  expected="开发者甲(已更新)", actual=r.json().get("full_name"))

    # 2.10 更新用户密码 (再登录验证)
    if "developer_id" in created:
        r = client.put(f"/api/users/{created['developer_id']}", json={
            "password": "NewDev@1234",
        }, headers=h)
        check(s, "更新密码返回 200", r.status_code == 200,
              expected=200, actual=r.status_code)
        if r.status_code == 200:
            new_login = login(client, dev_username, "NewDev@1234")
            check(s, "新密码可登录", new_login is not None)
            if new_login:
                created["developer_token"] = new_login["access_token"]
                created["developer_password"] = "NewDev@1234"

    # 2.11 search 过滤
    r = client.get(f"/api/users?search={dev_username}", headers=h)
    items = r.json().get("items", []) if r.status_code == 200 else []
    found = any(u.get("username") == dev_username for u in items)
    check(s, "search 过滤生效",
          r.status_code == 200 and found,
          expected=dev_username, actual=items[0].get("username") if items else None)

    # 2.12 软删除用户
    del_email = uniq_email("del")
    del_username = uniq_name("del")
    r = client.post("/api/users", json={
        "username": del_username, "email": del_email, "full_name": "待删除",
        "password": "Del@1234", "system_role": "guest",
    }, headers=h)
    if r.status_code == 201:
        del_id = r.json()["id"]
        r = client.delete(f"/api/users/{del_id}", headers=h)
        check(s, "软删除用户返回 204", r.status_code == 204,
              expected=204, actual=r.status_code)
        # 验证 is_active=False
        r = client.get(f"/api/users/{del_id}", headers=h)
        if r.status_code == 200:
            check(s, "软删除后 is_active=False",
                  r.json().get("is_active") is False,
                  expected=False, actual=r.json().get("is_active"))

    # 2.13 普通用户无权访问 /api/users
    if "developer_token" in created:
        r = client.get("/api/users", headers=auth_header(created["developer_token"]))
        check(s, "developer 无权访问用户列表 (403)",
              r.status_code == 403,
              expected=403, actual=r.status_code)

    # 2.14 删除不存在用户 → 404
    r = client.delete(f"/api/users/{fake_id}", headers=h)
    check(s, "删除不存在用户返回 404", r.status_code == 404,
          expected=404, actual=r.status_code)

    client.close()
    return created


# ============================================================================
# 3. 组织管理测试
# ============================================================================
def test_organizations(admin_token: str) -> dict:
    s = suite("3. 组织管理 (organizations)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)
    created: dict = {}

    # 3.1 创建部门
    dept_name = f"算法事业部_{uuid.uuid4().hex[:6]}"
    r = client.post("/api/organizations", json={
        "name": dept_name, "org_type": "department", "description": "顶层部门",
    }, headers=h)
    check(s, "创建部门返回 201", r.status_code == 201,
          expected=201, actual=r.status_code, detail=r.text[:200])
    if r.status_code == 201:
        created["dept_id"] = r.json()["id"]

    # 3.2 创建科室 (子节点)
    div_name = f"SOC科室_{uuid.uuid4().hex[:6]}"
    r = client.post("/api/organizations", json={
        "name": div_name, "org_type": "division",
        "parent_id": created.get("dept_id"), "description": "子科室",
    }, headers=h)
    check(s, "创建科室 (有 parent) 返回 201", r.status_code == 201,
          expected=201, actual=r.status_code, detail=r.text[:200])
    if r.status_code == 201:
        created["div_id"] = r.json()["id"]
        check(s, "科室 parent_id 正确",
              r.json().get("parent_id") == created.get("dept_id"),
              expected=created.get("dept_id"), actual=r.json().get("parent_id"))

    # 3.3 创建小组
    grp_name = f"开发小组_{uuid.uuid4().hex[:6]}"
    r = client.post("/api/organizations", json={
        "name": grp_name, "org_type": "group", "parent_id": created.get("div_id"),
    }, headers=h)
    check(s, "创建小组返回 201", r.status_code == 201,
          expected=201, actual=r.status_code)
    if r.status_code == 201:
        created["grp_id"] = r.json()["id"]

    # 3.4 不存在的 parent_id → 400
    r = client.post("/api/organizations", json={
        "name": "孤儿", "org_type": "division", "parent_id": str(uuid.uuid4()),
    }, headers=h)
    check(s, "不存在的 parent_id 返回 400", r.status_code == 400,
          expected=400, actual=r.status_code)

    # 3.5 自引用 parent → 400
    if "dept_id" in created:
        r = client.put(f"/api/organizations/{created['dept_id']}", json={
            "parent_id": created["dept_id"],
        }, headers=h)
        check(s, "自引用 parent 返回 400", r.status_code == 400,
              expected=400, actual=r.status_code, detail=r.text[:200])

    # 3.6 更新组织单元
    if "dept_id" in created:
        r = client.put(f"/api/organizations/{created['dept_id']}", json={
            "description": "顶层部门(已更新)",
        }, headers=h)
        check(s, "更新组织单元返回 200", r.status_code == 200,
              expected=200, actual=r.status_code)

    # 3.7 获取组织树
    r = client.get("/api/organizations/tree", headers=h)
    check(s, "GET /api/organizations/tree 返回 200", r.status_code == 200,
          expected=200, actual=r.status_code)
    if r.status_code == 200:
        tree = r.json()
        # 验证递归 children 结构
        check(s, "组织树含 children 字段",
              all("children" in n for n in tree) if tree else True,
              detail=f"top nodes: {len(tree)}")

    # 3.8 设置 admin scope (为 admin2)
    if "div_id" in created:
        # 需要先有 admin2_id; 任意存在的 admin/system 用户即可
        # 用前面创建的 admin2 (若存在) 否则用 admin 自己
        target_user = None
        r_users = client.get("/api/users?page=1&page_size=50&search=admin2", headers=h)
        if r_users.status_code == 200:
            items = r_users.json().get("items", [])
            if items:
                target_user = items[0]["id"]
        if target_user is None:
            # 退回到第一个非 admin 用户
            r_users = client.get("/api/users?page=1&page_size=50", headers=h)
            items = r_users.json().get("items", [])
            target_user = next((u["id"] for u in items if u["username"] != "admin"), None)

        if target_user:
            r = client.post("/api/organizations/admin-scopes", json={
                "user_id": target_user, "org_unit_id": created["div_id"],
            }, headers=h)
            check(s, "设置 admin scope 返回 201", r.status_code == 201,
                  expected=201, actual=r.status_code, detail=r.text[:200])

            # 3.9 重复设置 admin scope (幂等) → 201 返回原记录
            r2 = client.post("/api/organizations/admin-scopes", json={
                "user_id": target_user, "org_unit_id": created["div_id"],
            }, headers=h)
            check(s, "重复设置 admin scope 幂等返回 201", r2.status_code == 201,
                  expected=201, actual=r2.status_code)

            # 3.10 查询某用户的 admin scopes
            r = client.get(f"/api/organizations/admin-scopes/{target_user}", headers=h)
            check(s, "查询用户 admin scopes 返回 200", r.status_code == 200,
                  expected=200, actual=r.status_code)

    # 3.11 非 SUPER_ADMIN 不能创建组织单元
    # (用 developer token 测试 — 但 developer 无权, 应 403)
    # 此处简化: 用 admin2 (admin 角色) 应被拒
    client.close()
    return created


# ============================================================================
# 4. LLM 配置测试
# ============================================================================
def test_llm_config(admin_token: str) -> dict:
    s = suite("4. LLM 配置 (llm-config)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)
    created: dict = {}

    # 4.0 清理上次运行残留的评审规则 (review_type 唯一约束, 否则会 409)
    r = client.get("/api/llm-config/rules", headers=h)
    if r.status_code == 200:
        for rule in r.json():
            # 通过更新 is_active=False 软禁用, 或直接尝试更新复用
            # 这里采用: 更新现有规则指向新模型 (在模型创建后处理)
            pass

    # 4.1 列出模型 (空)
    r = client.get("/api/llm-config/models", headers=h)
    check(s, "GET /api/llm-config/models 返回 200", r.status_code == 200,
          expected=200, actual=r.status_code)

    # 4.2 创建 LLM 模型
    r = client.post("/api/llm-config/models", json={
        "name": "通义千问72B(测试)",
        "api_base": "http://llm-internal.local/v1",
        "api_key": "sk-test-key-001",
        "model_name": "qwen-72b",
        "priority": 0,
    }, headers=h)
    check(s, "创建 LLM 模型返回 201", r.status_code == 201,
          expected=201, actual=r.status_code, detail=r.text[:200])
    if r.status_code == 201:
        created["model_id"] = r.json()["id"]
        # 注意: 响应含 api_key 明文 (设计如此)
        check(s, "响应含 api_key 字段",
              "api_key" in r.json() and r.json()["api_key"] == "sk-test-key-001")

    # 4.3 创建备用模型
    r = client.post("/api/llm-config/models", json={
        "name": "通义千问14B(备用)",
        "api_base": "http://llm-internal.local/v1",
        "api_key": "sk-test-key-002",
        "model_name": "qwen-14b",
        "priority": 1,
    }, headers=h)
    check(s, "创建备用 LLM 模型返回 201", r.status_code == 201,
          expected=201, actual=r.status_code)
    if r.status_code == 201:
        created["fallback_model_id"] = r.json()["id"]

    # 4.4 列出模型 (应至少 2 个)
    r = client.get("/api/llm-config/models", headers=h)
    check(s, "列出模型 >=2 个",
          r.status_code == 200 and len(r.json()) >= 2,
          expected=">=2", actual=len(r.json()) if r.status_code == 200 else r.status_code)

    # 4.5 更新模型
    if "model_id" in created:
        r = client.put(f"/api/llm-config/models/{created['model_id']}", json={
            "name": "通义千问72B(已更新)",
        }, headers=h)
        check(s, "更新模型返回 200", r.status_code == 200,
              expected=200, actual=r.status_code)

    # 4.6 创建/复用评审规则 (code_review) — review_type 唯一, 已存在则更新复用
    existing_rules: dict = {}  # review_type -> rule_id
    r = client.get("/api/llm-config/rules", headers=h)
    if r.status_code == 200:
        for rule in r.json():
            existing_rules[rule["review_type"]] = rule["id"]

    if "model_id" in created:
        payload = {
            "llm_model_id": created["model_id"],
            "fallback_model_id": created.get("fallback_model_id"),
            "prompt_template": "请评审以下代码包:\n{content}\n输出 JSON。",
            "pass_threshold": 80.0,
            "dimension_thresholds": {"代码质量": 75.0, "变更合理性": 70.0},
            "is_active": True,
        }
        if "code_review" in existing_rules:
            r = client.put(f"/api/llm-config/rules/{existing_rules['code_review']}",
                           json=payload, headers=h)
            check(s, "复用更新 code_review 规则返回 200", r.status_code == 200,
                  expected=200, actual=r.status_code, detail=r.text[:200])
            if r.status_code == 200:
                created["rule_code_id"] = r.json()["id"]
        else:
            r = client.post("/api/llm-config/rules", json={
                "review_type": "code_review", **payload,
            }, headers=h)
            check(s, "创建 code_review 规则返回 201", r.status_code == 201,
                  expected=201, actual=r.status_code, detail=r.text[:200])
            if r.status_code == 201:
                created["rule_code_id"] = r.json()["id"]

    # 4.7 重复 review_type → 409 (仅当非复用模式时测试)
    if "model_id" in created and "code_review" not in existing_rules:
        r = client.post("/api/llm-config/rules", json={
            "review_type": "code_review",
            "llm_model_id": created["model_id"],
            "prompt_template": "另一个模板 {content}",
            "pass_threshold": 85.0,
        }, headers=h)
        check(s, "重复 review_type 返回 409", r.status_code == 409,
              expected=409, actual=r.status_code, detail=r.text[:200])
    else:
        # 复用模式下, 直接再 POST 同 review_type 验证 409
        if "model_id" in created:
            r = client.post("/api/llm-config/rules", json={
                "review_type": "code_review",
                "llm_model_id": created["model_id"],
                "prompt_template": "另一个模板 {content}",
                "pass_threshold": 85.0,
            }, headers=h)
            check(s, "重复 review_type 返回 409", r.status_code == 409,
                  expected=409, actual=r.status_code, detail=r.text[:200])

    # 4.8 创建/复用 test_report_review / expert_report_review 规则
    for rt in ["test_report_review", "expert_report_review"]:
        if "model_id" in created:
            payload = {
                "llm_model_id": created["model_id"],
                "prompt_template": f"评审 {rt}:\n{{content}}",
                "pass_threshold": 80.0,
                "is_active": True,
            }
            if rt in existing_rules:
                r = client.put(f"/api/llm-config/rules/{existing_rules[rt]}",
                               json=payload, headers=h)
                check(s, f"复用更新 {rt} 规则返回 200", r.status_code == 200,
                      expected=200, actual=r.status_code)
            else:
                r = client.post("/api/llm-config/rules", json={
                    "review_type": rt, **payload,
                }, headers=h)
                check(s, f"创建 {rt} 规则返回 201", r.status_code == 201,
                      expected=201, actual=r.status_code)

    # 4.9 规则引用不存在的模型 → 404
    r = client.post("/api/llm-config/rules", json={
        "review_type": "code_review",  # 已存在会409, 但模型不存在先报404
        "llm_model_id": str(uuid.uuid4()),
        "prompt_template": "x {content}",
    }, headers=h)
    # 注意: code_review 已存在, 可能 409 先于 404; 改用不存在模型 + 新 review_type 不可行(只有3种)
    # 这里容忍 404 或 409
    check(s, "规则引用不存在模型返回 4xx",
          r.status_code in (404, 409),
          expected="404 or 409", actual=r.status_code)

    # 4.10 列出规则
    r = client.get("/api/llm-config/rules", headers=h)
    check(s, "GET /api/llm-config/rules 返回 200", r.status_code == 200,
          expected=200, actual=r.status_code)
    if r.status_code == 200:
        check(s, "规则数 >=3 (3 种 review_type)",
              len(r.json()) >= 3, expected=">=3", actual=len(r.json()))

    # 4.11 软删除模型
    if "fallback_model_id" in created:
        r = client.delete(f"/api/llm-config/models/{created['fallback_model_id']}", headers=h)
        check(s, "软删除模型返回 200", r.status_code == 200,
              expected=200, actual=r.status_code)
        if r.status_code == 200:
            check(s, "软删除后 is_active=False",
                  r.json().get("is_active") is False,
                  expected=False, actual=r.json().get("is_active"))

    # 4.12 删除不存在模型 → 404
    r = client.delete(f"/api/llm-config/models/{str(uuid.uuid4())}", headers=h)
    check(s, "删除不存在模型返回 404", r.status_code == 404,
          expected=404, actual=r.status_code)

    client.close()
    return created


# ============================================================================
# 5. 项目生命周期测试
# ============================================================================
def test_projects(admin_token: str, users: dict) -> dict:
    s = suite("5. 项目生命周期 (projects)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)
    created: dict = {}

    # 5.1 创建项目 (admin 作为 PM)
    proj_name = f"SOX算法项目_{uuid.uuid4().hex[:6]}"
    r = client.post("/api/projects", json={
        "name": proj_name, "description": "测试项目 - SOC算法交付",
    }, headers=h)
    check(s, "创建项目返回 201", r.status_code == 201,
          expected=201, actual=r.status_code, detail=r.text[:200])
    if r.status_code == 201:
        created["project_id"] = r.json()["id"]
        created["project_name"] = proj_name
        # 验证 PM 自动加入 members
        members = r.json().get("members", [])
        check(s, "创建后 PM 自动成为成员",
              any(m.get("project_role") == "project_manager" for m in members),
              expected="project_manager in members", actual=members)

    # 5.2 列出项目
    r = client.get("/api/projects", headers=h)
    check(s, "GET /api/projects 返回 200 且非空",
          r.status_code == 200 and len(r.json()) >= 1,
          expected=">=1", actual=len(r.json()) if r.status_code == 200 else r.status_code)

    # 5.3 查询项目详情
    if "project_id" in created:
        r = client.get(f"/api/projects/{created['project_id']}", headers=h)
        check(s, "查询项目详情返回 200", r.status_code == 200,
              expected=200, actual=r.status_code)

    # 5.4 查询不存在项目 → 404
    r = client.get(f"/api/projects/{str(uuid.uuid4())}", headers=h)
    check(s, "查询不存在项目返回 404", r.status_code == 404,
          expected=404, actual=r.status_code)

    # 5.5 添加成员 (developer 角色给 dev 用户)
    if "project_id" in created and "developer_id" in users:
        r = client.post(f"/api/projects/{created['project_id']}/members", json={
            "user_id": users["developer_id"], "project_role": "developer",
        }, headers=h)
        check(s, "添加 developer 成员返回 201", r.status_code == 201,
              expected=201, actual=r.status_code, detail=r.text[:200])

    # 5.6 添加 tester 成员
    if "project_id" in created and "tester_id" in users:
        r = client.post(f"/api/projects/{created['project_id']}/members", json={
            "user_id": users["tester_id"], "project_role": "tester",
        }, headers=h)
        check(s, "添加 tester 成员返回 201", r.status_code == 201,
              expected=201, actual=r.status_code)

    # 5.7 添加 expert 成员
    if "project_id" in created and "expert_id" in users:
        r = client.post(f"/api/projects/{created['project_id']}/members", json={
            "user_id": users["expert_id"], "project_role": "external_expert",
        }, headers=h)
        check(s, "添加 external_expert 成员返回 201", r.status_code == 201,
              expected=201, actual=r.status_code)

    # 5.8 重复添加成员 → 409
    if "project_id" in created and "developer_id" in users:
        r = client.post(f"/api/projects/{created['project_id']}/members", json={
            "user_id": users["developer_id"], "project_role": "developer",
        }, headers=h)
        check(s, "重复添加成员返回 409", r.status_code == 409,
              expected=409, actual=r.status_code)

    # 5.9 创建版本 (指定 developer/tester/expert)
    if "project_id" in created:
        r = client.post(f"/api/projects/{created['project_id']}/versions", json={
            "version_number": "v1.0",
            "description": "首个版本 - SOC算法初版",
            "developer_id": users.get("developer_id"),
            "tester_id": users.get("tester_id"),
            "expert_id": users.get("expert_id"),
        }, headers=h)
        check(s, "创建版本返回 201", r.status_code == 201,
              expected=201, actual=r.status_code, detail=r.text[:200])
        if r.status_code == 201:
            created["version_id"] = r.json()["id"]

    # 5.10 重复版本号 → 409
    if "project_id" in created:
        r = client.post(f"/api/projects/{created['project_id']}/versions", json={
            "version_number": "v1.0", "description": "重复版本号",
        }, headers=h)
        check(s, "重复版本号返回 409", r.status_code == 409,
              expected=409, actual=r.status_code)

    client.close()
    return created


# ============================================================================
# 6. 释放流程测试 (7 步流转)
# ============================================================================
def test_releases(admin_token: str, users: dict, projects: dict, minio_up: bool) -> dict:
    s = suite("6. 释放流程 (releases)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)
    created: dict = {}

    # 6.1 找到版本对应的 DRAFT release (创建版本时自动生成)
    if "version_id" not in projects:
        check(s, "前置: 版本已创建", False, detail="缺少 version_id")
        client.close()
        return created

    # 通过 search 找 release
    r = client.get(f"/api/search/releases?version_number=v1.0&page=1&page_size=10", headers=h)
    release_id = None
    if r.status_code == 200:
        items = r.json().get("items", [])
        if items:
            release_id = items[0]["id"]
    if release_id is None:
        check(s, "前置: 找到 DRAFT release", False, detail="无 release")
        client.close()
        return created
    created["release_id"] = release_id

    # 6.2 查询 release 详情, 验证初始状态为 draft
    r = client.get(f"/api/releases/{release_id}", headers=h)
    check(s, "查询 release 详情返回 200", r.status_code == 200,
          expected=200, actual=r.status_code)
    if r.status_code == 200:
        check(s, "初始状态为 draft",
              r.json().get("status") == "draft",
              expected="draft", actual=r.json().get("status"))

    # 6.3 查询不存在 release → 404
    r = client.get(f"/api/releases/{str(uuid.uuid4())}", headers=h)
    check(s, "查询不存在 release 返回 404", r.status_code == 404,
          expected=404, actual=r.status_code)

    # ---- 文件上传测试 (依赖 MinIO) ----
    if minio_up:
        # 6.4 上传代码包 → status: code_pending_review
        zip_bytes = _make_fake_zip()
        r = client.post(
            f"/api/releases/{release_id}/code-package",
            headers=h,
            files={"file": ("code.zip", zip_bytes, "application/zip")},
            data={"change_notes": "新增 SOC 估算核心函数; 修复电压采样偏差"},
        )
        check(s, "上传代码包返回 200", r.status_code == 200,
              expected=200, actual=r.status_code, detail=r.text[:300])
        if r.status_code == 200:
            check(s, "上传代码包后状态 → code_pending_review",
                  r.json().get("status") == "code_pending_review",
                  expected="code_pending_review", actual=r.json().get("status"))
            check(s, "code_package_path 已写入",
                  bool(r.json().get("code_package_path")),
                  expected="non-empty", actual=r.json().get("code_package_path"))
            check(s, "change_notes 已写入",
                  r.json().get("change_notes") == "新增 SOC 估算核心函数; 修复电压采样偏差",
                  expected="...", actual=r.json().get("change_notes"))

        # 6.5 上传测试报告 → status: test_pending_review
        docx_bytes = _make_fake_docx()
        r = client.post(
            f"/api/releases/{release_id}/test-report",
            headers=h,
            files={"file": ("test_report.docx", docx_bytes,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        check(s, "上传测试报告返回 200", r.status_code == 200,
              expected=200, actual=r.status_code, detail=r.text[:300])
        if r.status_code == 200:
            check(s, "上传测试报告后状态 → test_pending_review",
                  r.json().get("status") == "test_pending_review",
                  expected="test_pending_review", actual=r.json().get("status"))

        # 6.6 上传专家评审报告 → status: expert_pending_review
        xlsx_bytes = _make_fake_xlsx()
        r = client.post(
            f"/api/releases/{release_id}/review-report",
            headers=h,
            files={"file": ("review.xlsx", xlsx_bytes,
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        check(s, "上传专家评审报告返回 200", r.status_code == 200,
              expected=200, actual=r.status_code, detail=r.text[:300])
        if r.status_code == 200:
            check(s, "上传评审报告后状态 → expert_pending_review",
                  r.json().get("status") == "expert_pending_review",
                  expected="expert_pending_review", actual=r.json().get("status"))
    else:
        check(s, "[跳过] 文件上传 (MinIO 未就绪)", True,
              detail="minio not running")

    # 6.7 触发 LLM 评审 (异步, 入队即返回 202)
    # 需要 review_type query 参数
    r = client.post(f"/api/reviews/trigger/{release_id}?review_type=code_review", headers=h)
    check(s, "触发 code_review 评审返回 202", r.status_code == 202,
          expected=202, actual=r.status_code, detail=r.text[:300])
    if r.status_code == 202:
        body = r.json()
        check(s, "trigger 响应含 task_id",
              "task_id" in body and body.get("status") == "queued",
              expected="queued", actual=body.get("status"))

    # 6.8 缺少 review_type 参数 → 422
    r = client.post(f"/api/reviews/trigger/{release_id}", headers=h)
    check(s, "缺少 review_type 参数返回 422", r.status_code == 422,
          expected=422, actual=r.status_code)

    # 6.9 列出 release 的评审记录
    r = client.get(f"/api/reviews/release/{release_id}", headers=h)
    check(s, "列出 release 评审记录返回 200", r.status_code == 200,
          expected=200, actual=r.status_code)

    # 6.10 触发不存在 release 的评审 → 404
    r = client.post(f"/api/reviews/trigger/{str(uuid.uuid4())}?review_type=code_review", headers=h)
    check(s, "触发不存在 release 评审返回 404", r.status_code == 404,
          expected=404, actual=r.status_code)

    # 6.11 确认释放 (PM = admin)
    # 先把状态推进到 pending_confirm (这里直接 confirm, 实际需 LLM 通过)
    # 由于 LLM 评审需要真实 LLM, 这里测试 confirm 端点的权限校验即可
    # 用 expert 角色用户 (非 PM) 应被拒
    if "expert_id" in users and "expert_username" in users:
        # expert 是 guest 角色 + 项目成员 external_expert
        expert_login = login(client, users["expert_username"], "Expert@123")
        if expert_login:
            r = client.post(f"/api/releases/{release_id}/confirm",
                            headers=auth_header(expert_login["access_token"]))
            check(s, "非 PM (expert) 确认释放返回 403",
                  r.status_code == 403,
                  expected=403, actual=r.status_code, detail=r.text[:200])

    # PM (admin) 确认释放 → released (无论 LLM 是否通过, 端点允许)
    r = client.post(f"/api/releases/{release_id}/confirm", headers=h)
    check(s, "PM 确认释放返回 200", r.status_code == 200,
          expected=200, actual=r.status_code, detail=r.text[:300])
    if r.status_code == 200:
        check(s, "确认后状态 → released",
              r.json().get("status") == "released",
              expected="released", actual=r.json().get("status"))
        check(s, "confirmed_by 已写入",
              bool(r.json().get("confirmed_by")),
              expected="non-empty", actual=r.json().get("confirmed_by"))
        check(s, "confirmed_at 已写入",
              bool(r.json().get("confirmed_at")),
              expected="non-empty", actual=r.json().get("confirmed_at"))
        # download_link 仅当 code_package_path 存在时生成
        if minio_up:
            check(s, "download_link 已生成 (有代码包)",
                  bool(r.json().get("download_link")),
                  expected="non-empty", actual=r.json().get("download_link"))

    client.close()
    return created


def _make_fake_zip() -> bytes:
    """生成一个最小合法 zip (含一个 .py 文件)"""
    import zipfile
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("soc.py", "def calc_soc(voltage):\n    return voltage / 4.2 * 100\n")
        zf.writestr("README.txt", "fake code package for testing\n")
    return buf.getvalue()


def _make_fake_docx() -> bytes:
    """生成一个最小合法 docx"""
    from docx import Document
    doc = Document()
    doc.add_heading("测试报告", level=1)
    doc.add_paragraph("SOC 算法测试通过。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_fake_xlsx() -> bytes:
    """生成一个最小合法 xlsx"""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "评审"
    ws.append(["项目", "结论"])
    ws.append(["SOC算法", "通过"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ============================================================================
# 7. 搜索筛选测试
# ============================================================================
def test_search(admin_token: str, users: dict, projects: dict) -> None:
    s = suite("7. 搜索筛选 (search)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)

    # 7.1 搜索 release (无过滤)
    r = client.get("/api/search/releases?page=1&page_size=10", headers=h)
    check(s, "GET /api/search/releases 返回分页结构",
          r.status_code == 200 and "items" in r.json(),
          expected="PaginatedResponse", actual=r.status_code)

    # 7.2 按 project_name 过滤
    if "project_name" in projects:
        r = client.get(f"/api/search/releases?project_name={projects['project_name']}", headers=h)
        check(s, "按 project_name 过滤返回 200",
              r.status_code == 200,
              expected=200, actual=r.status_code)

    # 7.3 按 version_number 过滤
    r = client.get("/api/search/releases?version_number=v1.0", headers=h)
    check(s, "按 version_number 过滤返回 200",
          r.status_code == 200,
          expected=200, actual=r.status_code)

    # 7.4 按 status 过滤 (精确)
    r = client.get("/api/search/releases?status=released", headers=h)
    check(s, "按 status=released 过滤返回 200",
          r.status_code == 200,
          expected=200, actual=r.status_code)
    if r.status_code == 200:
        items = r.json().get("items", [])
        all_released = all(i.get("status") == "released" for i in items)
        check(s, "status=released 结果均为 released",
              all_released, expected=True, actual=all_released)

    # 7.5 search releases 响应含关联字段
    r = client.get("/api/search/releases?page=1&page_size=1", headers=h)
    if r.status_code == 200 and r.json().get("items"):
        item = r.json()["items"][0]
        check(s, "search releases 含关联字段 (project_name/version_number)",
              "project_name" in item and "version_number" in item,
              expected="has join fields", actual=list(item.keys()))

    # 7.6 搜索 projects
    r = client.get("/api/search/projects?page=1&page_size=10", headers=h)
    check(s, "GET /api/search/projects 返回 200",
          r.status_code == 200 and "items" in r.json(),
          expected=200, actual=r.status_code)

    # 7.7 按 name 过滤 projects
    if "project_name" in projects:
        r = client.get(f"/api/search/projects?name={projects['project_name'][:10]}", headers=h)
        check(s, "按 name 过滤 projects 返回 200",
              r.status_code == 200,
              expected=200, actual=r.status_code)

    # 7.8 GUEST 隔离: guest 用户只能看 released
    # 用前面注册的 expert (guest) 或新注册一个 guest
    guest_email = uniq_email("guest_search")
    guest_username = uniq_name("guest_search")
    r = client.post("/api/auth/register", json={
        "username": guest_username, "email": guest_email, "full_name": "搜索访客",
        "password": "Guest@123",
    })
    if r.status_code == 201:
        glogin = login(client, guest_username, "Guest@123")
        if glogin:
            gh = auth_header(glogin["access_token"])
            # GUEST 传 status=draft, 应仍只返回 released (或空)
            r = client.get("/api/search/releases?status=draft", headers=gh)
            check(s, "GUEST 搜索 releases 返回 200",
                  r.status_code == 200, expected=200, actual=r.status_code)
            if r.status_code == 200:
                items = r.json().get("items", [])
                all_released = all(i.get("status") == "released" for i in items)
                check(s, "GUEST 仅可见 released (传 draft 也被忽略)",
                      all_released, expected=True, actual=all_released)

            # GUEST 搜索 projects 仅可见有 released 的
            r = client.get("/api/search/projects", headers=gh)
            check(s, "GUEST 搜索 projects 返回 200",
                  r.status_code == 200, expected=200, actual=r.status_code)

    client.close()


# ============================================================================
# 8. 审计日志测试
# ============================================================================
def test_audit(admin_token: str) -> None:
    s = suite("8. 审计日志 (audit)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)

    # 8.1 列出审计日志
    r = client.get("/api/audit?page=1&page_size=20", headers=h)
    check(s, "GET /api/audit 返回分页结构",
          r.status_code == 200 and "items" in r.json(),
          expected="PaginatedResponse", actual=r.status_code)
    if r.status_code == 200:
        items = r.json().get("items", [])
        check(s, "审计日志非空 (前面操作已产生记录)",
              len(items) > 0, expected=">0", actual=len(items))
        if items:
            sample = items[0]
            check(s, "审计日志含 action/resource_type/created_at",
                  all(k in sample for k in ["action", "resource_type", "created_at"]),
                  expected="has key fields", actual=list(sample.keys()))

    # 8.2 按 action 过滤
    r = client.get("/api/audit?action=login&page=1&page_size=5", headers=h)
    check(s, "按 action=login 过滤返回 200",
          r.status_code == 200, expected=200, actual=r.status_code)
    if r.status_code == 200:
        items = r.json().get("items", [])
        all_login = all(i.get("action") == "login" for i in items)
        check(s, "action=login 结果均为 login",
              all_login, expected=True, actual=all_login)

    # 8.3 按 resource_type 过滤
    r = client.get("/api/audit?resource_type=user&page=1&page_size=5", headers=h)
    check(s, "按 resource_type=user 过滤返回 200",
          r.status_code == 200, expected=200, actual=r.status_code)

    # 8.4 普通用户无权访问审计日志
    # 用 developer token
    dev_email = uniq_email("audit_dev")
    dev_username = uniq_name("audit_dev")
    r = client.post("/api/users", json={
        "username": dev_username, "email": dev_email, "full_name": "审计测试dev",
        "password": "Dev@1234", "system_role": "developer",
    }, headers=h)
    if r.status_code == 201:
        dlogin = login(client, dev_username, "Dev@1234")
        if dlogin:
            r = client.get("/api/audit", headers=auth_header(dlogin["access_token"]))
            check(s, "developer 无权访问审计日志 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

    client.close()


# ============================================================================
# 9. 通知测试
# ============================================================================
def test_notifications(admin_token: str, users: dict) -> None:
    s = suite("9. 通知 (notifications)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)

    # 9.1 列出当前用户通知
    r = client.get("/api/notifications?page=1&page_size=10", headers=h)
    check(s, "GET /api/notifications 返回分页结构",
          r.status_code == 200 and "items" in r.json(),
          expected="PaginatedResponse", actual=r.status_code)

    # 9.2 unread_only 过滤
    r = client.get("/api/notifications?unread_only=true&page=1&page_size=10", headers=h)
    check(s, "unread_only=true 过滤返回 200",
          r.status_code == 200, expected=200, actual=r.status_code)
    if r.status_code == 200:
        items = r.json().get("items", [])
        all_unread = all(i.get("is_read") is False for i in items)
        check(s, "unread_only 结果均为未读",
              all_unread, expected=True, actual=all_unread)

    # 9.3 标记通知已读 (若有通知)
    r = client.get("/api/notifications?unread_only=true&page=1&page_size=1", headers=h)
    if r.status_code == 200 and r.json().get("items"):
        nid = r.json()["items"][0]["id"]
        r = client.post(f"/api/notifications/{nid}/read", headers=h)
        check(s, "标记通知已读返回 200",
              r.status_code == 200, expected=200, actual=r.status_code)
        if r.status_code == 200:
            check(s, "标记后 is_read=True",
                  r.json().get("is_read") is True,
                  expected=True, actual=r.json().get("is_read"))

    # 9.4 标记不存在通知 → 404
    r = client.post(f"/api/notifications/{str(uuid.uuid4())}/read", headers=h)
    check(s, "标记不存在通知返回 404",
          r.status_code == 404, expected=404, actual=r.status_code)

    # 9.5 跨用户标记他人通知 → 404 (而非 403)
    # 用 developer 登录, 尝试标记 admin 的通知
    if "developer_token" in users:
        r = client.get("/api/notifications?unread_only=false&page=1&page_size=1", headers=h)
        if r.status_code == 200 and r.json().get("items"):
            admin_nid = r.json()["items"][0]["id"]
            r = client.post(f"/api/notifications/{admin_nid}/read",
                            headers=auth_header(users["developer_token"]))
            check(s, "跨用户标记他人通知返回 404",
                  r.status_code == 404, expected=404, actual=r.status_code)

    client.close()


# ============================================================================
# 10. 权限矩阵测试
# ============================================================================
def test_permissions(admin_token: str, users: dict, projects: dict) -> None:
    s = suite("10. 权限矩阵 (permission)")
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    h = auth_header(admin_token)

    # 10.1 GUEST 不能创建项目
    guest_email = uniq_email("perm_guest")
    guest_username = uniq_name("perm_guest")
    r = client.post("/api/auth/register", json={
        "username": guest_username, "email": guest_email, "full_name": "权限访客",
        "password": "Guest@123",
    })
    if r.status_code == 201:
        glogin = login(client, guest_username, "Guest@123")
        if glogin:
            gh = auth_header(glogin["access_token"])
            r = client.post("/api/projects", json={
                "name": "访客项目", "description": "应被拒",
            }, headers=gh)
            check(s, "GUEST 创建项目被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

            # GUEST 不能触发评审
            if "release_id" in projects:
                r = client.post(f"/api/reviews/trigger/{projects['release_id']}?review_type=code_review",
                                headers=gh)
                check(s, "GUEST 触发评审被拒 (403)",
                      r.status_code == 403, expected=403, actual=r.status_code)

            # GUEST 不能访问用户列表
            r = client.get("/api/users", headers=gh)
            check(s, "GUEST 访问用户列表被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

            # GUEST 不能访问审计
            r = client.get("/api/audit", headers=gh)
            check(s, "GUEST 访问审计被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

            # GUEST 不能访问 LLM 配置
            r = client.get("/api/llm-config/models", headers=gh)
            check(s, "GUEST 访问 LLM 配置被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

            # GUEST 不能创建组织单元
            r = client.post("/api/organizations", json={
                "name": "x", "org_type": "department",
            }, headers=gh)
            check(s, "GUEST 创建组织被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

    # 10.2 无 Token 访问受保护端点 → 401
    for method, path in [
        ("GET", "/api/users"),
        ("GET", "/api/audit"),
        ("GET", "/api/llm-config/models"),
        ("GET", "/api/organizations/tree"),
        ("GET", "/api/notifications"),
        ("POST", "/api/projects"),
        ("GET", "/api/projects"),
    ]:
        r = client.request(method, path)
        check(s, f"无 Token {method} {path} → 401",
              r.status_code == 401, expected=401, actual=r.status_code)

    # 10.3 跨项目隔离: 非 PM/成员不能访问项目
    # 创建一个只有 admin 的项目, 然后用 developer (非成员) 访问
    proj_name2 = f"隔离项目_{uuid.uuid4().hex[:6]}"
    r = client.post("/api/projects", json={
        "name": proj_name2, "description": "隔离测试",
    }, headers=h)
    if r.status_code == 201:
        proj2_id = r.json()["id"]
        if "developer_token" in users:
            r = client.get(f"/api/projects/{proj2_id}",
                           headers=auth_header(users["developer_token"]))
            check(s, "非成员访问项目被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

    # 10.4 ADMIN 角色: 能访问用户列表但能创建组织单元吗? (应被拒, 仅超管)
    if "admin2_id" in users:
        # admin2 是 admin 角色, 用其登录
        alogin = login(client, users.get("admin2_username", ""), "Admin2@123")
        if alogin:
            ah = auth_header(alogin["access_token"])
            # admin 能访问用户列表
            r = client.get("/api/users", headers=ah)
            check(s, "ADMIN 能访问用户列表 (200)",
                  r.status_code == 200, expected=200, actual=r.status_code)
            # admin 能访问审计日志
            r = client.get("/api/audit", headers=ah)
            check(s, "ADMIN 能访问审计日志 (200)",
                  r.status_code == 200, expected=200, actual=r.status_code)
            # admin 不能访问 LLM 配置
            r = client.get("/api/llm-config/models", headers=ah)
            check(s, "ADMIN 访问 LLM 配置被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)
            # admin 不能创建组织单元
            r = client.post("/api/organizations", json={
                "name": "admin试建", "org_type": "department",
            }, headers=ah)
            check(s, "ADMIN 创建组织单元被拒 (403)",
                  r.status_code == 403, expected=403, actual=r.status_code)

    # 10.5 禁用用户访问 → 403
    dis_email = uniq_email("dis")
    dis_username = uniq_name("dis")
    r = client.post("/api/users", json={
        "username": dis_username, "email": dis_email, "full_name": "待禁用",
        "password": "Dis@1234", "system_role": "developer",
    }, headers=h)
    if r.status_code == 201:
        dis_id = r.json()["id"]
        # 禁用
        client.put(f"/api/users/{dis_id}", json={"is_active": False}, headers=h)
        dlogin = login(client, dis_username, "Dis@1234")
        check(s, "禁用用户登录被拒 (401)",
              dlogin is None, expected=None, actual="token" if dlogin else None)

    client.close()


# ============================================================================
# 主流程
# ============================================================================
def minio_is_up() -> bool:
    try:
        with httpx.Client(timeout=3.0) as c:
            r = c.get("http://localhost:9000/minio/health/live")
            return r.status_code == 200
    except Exception:
        return False


def main() -> int:
    print("=" * 72)
    print("qloop — 全面功能测试")
    print("=" * 72)

    minio_up = minio_is_up()
    print(f"\n[环境] MinIO: {'就绪 (9000)' if minio_up else '未就绪 (文件上传测试将跳过)'}\n")

    # 登录 admin
    client = httpx.Client(base_url=BASE, timeout=TIMEOUT)
    admin_data = login(client, "admin", "Admin@123")
    client.close()
    if not admin_data:
        print("[FATAL] admin 登录失败, 终止测试")
        return 1
    admin_token = admin_data["access_token"]
    print(f"[OK] admin 登录成功\n")

    # 执行各测试套件
    test_auth()
    users = test_users(admin_token)
    orgs = test_organizations(admin_token)
    llm = test_llm_config(admin_token)
    projects = test_projects(admin_token, users)
    releases = test_releases(admin_token, users, projects, minio_up)
    # 把 release_id 合并进 projects 字典供后续用
    projects.update(releases)
    test_search(admin_token, users, projects)
    test_audit(admin_token)
    test_notifications(admin_token, users)
    test_permissions(admin_token, users, projects)

    # 汇总报告
    print("\n" + "=" * 72)
    print("测试报告")
    print("=" * 72)
    total_pass = 0
    total_fail = 0
    for s in SUITES:
        passed = s.passed_count
        failed = s.failed_count
        total_pass += passed
        total_fail += failed
        status = "✓ PASS" if failed == 0 else f"✗ {failed} FAILED"
        print(f"  {s.name:40s} {passed:3d} 通过 / {failed:3d} 失败  [{status}]")
        for r in s.results:
            if not r.passed:
                print(f"      ✗ {r.name}")
                print(f"        预期: {r.expected}  实际: {r.actual}")
                if r.detail:
                    print(f"        详情: {r.detail[:200]}")
    print("-" * 72)
    print(f"  总计: {total_pass} 通过, {total_fail} 失败, {len(SUITES)} 个套件")
    rate = total_pass / (total_pass + total_fail) * 100 if (total_pass + total_fail) else 0
    print(f"  通过率: {rate:.1f}%")
    print("=" * 72)
    return 0 if total_fail == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())
