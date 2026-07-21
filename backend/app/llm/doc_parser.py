"""Document parsing helpers for LLM review.

These helpers turn an uploaded artifact into plain text that can be
embedded into an LLM review prompt. Supported formats:

* ``.docx`` -> :func:`parse_docx` (via python-docx)
* ``.xlsx`` -> :func:`parse_xlsx` (via openpyxl)
* ``.zip`` -> :func:`parse_zip` (auto-extract and recursively parse inner
  documents; supports nested .docx/.xlsx/.txt/.md/.csv/.json/.yaml)
* ``.txt`` / ``.md`` / ``.csv`` / ``.json`` / ``.yaml`` / ``.ini`` /
  ``.log`` / ``.rst`` -> decoded as text (tries utf-8 -> gbk -> latin-1)
* ``.pdf`` -> descriptive message (not supported without pypdf)
* ``.doc`` / ``.xls`` -> descriptive message (legacy binary formats)

The legacy binary formats and any other extension return a descriptive
message instead of raising, so the review engine can still produce a
(failed) result.
"""

from __future__ import annotations

import io
import os
import zipfile
from typing import Dict, List

from docx import Document
from openpyxl import load_workbook


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
# Text file extensions that can be decoded directly.
_TEXT_EXTENSIONS: frozenset = frozenset({
    ".txt", ".md", ".markdown", ".csv", ".log", ".json",
    ".yaml", ".yml", ".ini", ".cfg", ".rst", ".xml", ".html", ".htm",
})

# Extensions that ``parse_zip`` will attempt to parse inside an archive.
_ZIP_PARSEABLE_EXTENSIONS: frozenset = frozenset({
    ".docx", ".xlsx",
}) | _TEXT_EXTENSIONS

# Maximum number of files to parse inside a ZIP archive (DoS guard).
_MAX_ZIP_ENTRIES = 50

# Maximum total text length returned (LLM prompt guard).
_MAX_TEXT_LENGTH = 100_000


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _truncate(text: str) -> str:
    """Truncate text to a safe length for LLM prompts."""
    if len(text) > _MAX_TEXT_LENGTH:
        return text[:_MAX_TEXT_LENGTH] + "\n\n[... 内容已截断 ...]"
    return text


def _decode_text(file_data: bytes) -> str:
    """Try common encodings in order; return the first successful decode."""
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb18030", "latin-1"):
        try:
            return file_data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return "[无法解码文本文件，尝试的所有编码均失败]"


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------
def parse_docx(file_data: bytes) -> str:
    """Extract text from a ``.docx`` file.

    Both paragraph text and table cell text are included, in document
    order. Empty paragraphs are skipped.

    Args:
        file_data: The raw bytes of the ``.docx`` file.

    Returns:
        The extracted plain text. If parsing fails, an error message is
        returned.
    """
    try:
        document = Document(io.BytesIO(file_data))
    except Exception as exc:  # noqa: BLE001
        return f"[解析 Word 文档失败: {exc}]"

    chunks: List[str] = []

    # Iterate over the document body in order, interleaving paragraphs and
    # tables as they appear.
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        chunks.append("--- 表格 ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
        chunks.append("--- 表格结束 ---")

    return "\n".join(chunks)


# ---------------------------------------------------------------------------
# Excel (.xlsx)
# ---------------------------------------------------------------------------
def parse_xlsx(file_data: bytes) -> str:
    """Extract text from an ``.xlsx`` file.

    Every worksheet is traversed row by row; non-empty rows are joined with
    ``|`` separators and prefixed with a sheet header.

    Args:
        file_data: The raw bytes of the ``.xlsx`` file.

    Returns:
        The extracted plain text. If parsing fails, an error message is
        returned.
    """
    try:
        workbook = load_workbook(io.BytesIO(file_data), data_only=True, read_only=True)
    except Exception as exc:  # noqa: BLE001
        return f"[解析 Excel 文档失败: {exc}]"

    chunks: List[str] = []
    for sheet in workbook.worksheets:
        chunks.append(f"=== Sheet: {sheet.title} ===")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            # Skip completely empty rows.
            if all(cell is None or str(cell).strip() == "" for cell in row):
                continue
            cells = [str(cell).strip() if cell is not None else "" for cell in row]
            chunks.append(" | ".join(cells))
            row_count += 1
        if row_count == 0:
            chunks.append("(空工作表)")
        chunks.append(f"=== Sheet 结束: {sheet.title} ===")

    try:
        workbook.close()
    except Exception:  # noqa: BLE001
        pass

    return "\n".join(chunks)


# ---------------------------------------------------------------------------
# ZIP archive
# ---------------------------------------------------------------------------
def _extract_zip(file_data: bytes) -> Dict[str, bytes]:
    """Extract a ZIP archive into ``{path: content}``.

    Skips directories and macOS metadata entries. Capped at
    ``_MAX_ZIP_ENTRIES`` to mitigate ZIP-bomb style inputs.
    """
    files: Dict[str, bytes] = {}
    with zipfile.ZipFile(io.BytesIO(file_data)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            if "__MACOSX" in info.filename or os.path.basename(info.filename).startswith("._"):
                continue
            try:
                files[info.filename] = zf.read(info)
            except (RuntimeError, zipfile.BadZipFile, OSError):
                continue
            if len(files) >= _MAX_ZIP_ENTRIES:
                break
    return files


def parse_zip(file_data: bytes) -> str:
    """Extract a ZIP archive and parse all inner documents.

    Each parseable file is rendered with a ``--- {path} ---`` header so the
    LLM can attribute content to specific files. Non-parseable files are
    listed but skipped (so the LLM still sees what was uploaded).

    Args:
        file_data: The raw bytes of a ``.zip`` archive.

    Returns:
        The concatenated plain text of all parseable inner files.
    """
    try:
        files = _extract_zip(file_data)
    except zipfile.BadZipFile as exc:
        return f"[解压 ZIP 失败: {exc}]"
    except Exception as exc:  # noqa: BLE001
        return f"[读取 ZIP 失败: {exc}]"

    if not files:
        return "[ZIP 包内未找到任何可解析的文件]"

    chunks: List[str] = []
    for path, content in files.items():
        ext = os.path.splitext(path)[1].lower()

        if ext in _ZIP_PARSEABLE_EXTENSIONS:
            inner_text = parse_document(content, path)
            chunks.append(f"--- {path} ---")
            chunks.append(inner_text)
            chunks.append("")  # blank line separator
        else:
            # Surface the file name so the LLM knows it existed but was
            # not parsed (e.g. .exe, .dll, .png).
            chunks.append(f"--- {path} ---")
            chunks.append(f"[跳过不支持的文件类型: {ext or '(无扩展名)'}]")
            chunks.append("")

    return _truncate("\n".join(chunks))


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
def parse_document(file_data: bytes, filename: str) -> str:
    """Parse a document based on its file extension.

    Supported extensions:
        * ``.docx`` -> :func:`parse_docx`
        * ``.xlsx`` -> :func:`parse_xlsx`
        * ``.zip`` -> :func:`parse_zip` (auto-extract + recursive parse)
        * ``.txt`` / ``.md`` / ``.csv`` / ``.json`` / ``.yaml`` / ``.ini`` /
          ``.log`` / ``.rst`` -> decoded as text
        * ``.pdf`` -> descriptive message (not supported)
        * ``.doc`` / ``.xls`` -> descriptive message (legacy binary formats)

    Any other extension falls back to a best-effort UTF-8 decode.

    Args:
        file_data: The raw bytes of the document.
        filename: The original file name (used to determine the parser).

    Returns:
        The extracted plain text, or a message explaining why parsing was
        skipped.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".docx":
        return parse_docx(file_data)
    if ext == ".xlsx":
        return parse_xlsx(file_data)
    if ext == ".zip":
        return parse_zip(file_data)
    if ext in _TEXT_EXTENSIONS:
        return _decode_text(file_data)
    if ext == ".pdf":
        return (
            "[当前版本暂不支持 PDF 直接解析。请将文档另存为 .docx 或导出为 "
            ".md/.txt 后重新上传；或将多个文档打包为 .zip（内含 .docx/.md/.txt "
            "等可解析文件）后上传。]"
        )
    if ext == ".doc":
        return (
            "[不支持旧版 .doc 格式，请将文档另存为 .docx 后重新上传以进行评审。]"
        )
    if ext == ".xls":
        return (
            "[不支持旧版 .xls 格式，请将文档另存为 .xlsx 后重新上传以进行评审。]"
        )

    # Last resort: try to decode as plain text.
    try:
        return file_data.decode("utf-8")
    except UnicodeDecodeError:
        return f"[无法解析的文档类型: {ext or '(无扩展名)'}]"
